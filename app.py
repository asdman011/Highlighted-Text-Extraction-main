from flask import Flask, render_template, request, send_file, session, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
import docx
import csv
import os
from datetime import datetime

app = Flask(__name__)
app.secret_key = "1111"  # Add a secret key for session management
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///highlighted_text.db'
db = SQLAlchemy(app)

# Define the Documents table
class Document(db.Model):
    doc_id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    highlights = db.relationship('Highlight', backref='document', lazy=True)

# Define the Highlights table
class Highlight(db.Model):
    highlight_id = db.Column(db.Integer, primary_key=True)
    doc_id = db.Column(db.Integer, db.ForeignKey('document.doc_id'), nullable=False)
    color = db.Column(db.String(50), nullable=False)
    text = db.Column(db.Text, nullable=False)

def get_color_name(highlight_color):
    # Check if the highlight_color is an instance of the enum
    if isinstance(highlight_color, docx.enum.text.WD_COLOR_INDEX):
        # Convert the enum value to a string
        return highlight_color.name
    else:
        return None

def extract_highlighted_text(docx_file, doc_id):
    # Open the input file as a docx document object
    doc = docx.Document(docx_file)
    # Loop through all the paragraphs in the document
    for para in doc.paragraphs:
        # Loop through all the runs in each paragraph
        for run in para.runs:
            # Check if the run has a highlight color
            highlight_color = run.font.highlight_color
            color_name = get_color_name(highlight_color)
            if color_name:
                # Add highlighted text to the database
                highlight = Highlight(doc_id=doc_id, color=color_name, text=run.text)
                db.session.add(highlight)
    db.session.commit()

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        docx_file = request.files['docxFile']
        if docx_file and docx_file.filename.endswith('.docx'):
            # Save the uploaded file to a temporary location
            temp_path = os.path.join('uploads', docx_file.filename)
            docx_file.save(temp_path)

            # Add document record to the database
            document = Document(filename=docx_file.filename)
            db.session.add(document)
            db.session.commit()

            # Process the uploaded file
            extract_highlighted_text(temp_path, document.doc_id)

            # Remove the temporary uploaded file
            os.remove(temp_path)

            return redirect(url_for('index'))

    documents = Document.query.all()
    return render_template('index.html', documents=documents)

@app.route('/highlights/<int:doc_id>')
def highlights(doc_id):
    document = Document.query.get_or_404(doc_id)
    highlights = Highlight.query.filter_by(doc_id=doc_id).all()
    highlighted_data = {}
    for highlight in highlights:
        highlighted_data.setdefault(highlight.color, []).append(highlight.text)
    return render_template('highlights.html', document=document, highlighted_data=highlighted_data)

@app.route('/download/<int:doc_id>')
def download_file(doc_id):
    # Open the output file as a csv writer object
    output_file = os.path.join('outputs', f'highlighted_text_{doc_id}.csv')
    with open(output_file, "w", newline="", encoding="utf-8") as csv_file:
        fieldnames = ['Color', 'Highlighted Text']
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames, quoting=csv.QUOTE_MINIMAL, lineterminator='\n')
        # Write header
        writer.writeheader()
        # Write rows in chunks
        chunk_size = 10000  # Set an appropriate chunk size
        highlights = Highlight.query.filter_by(doc_id=doc_id).all()
        for highlight in highlights:
            writer.writerow({'Color': highlight.color, 'Highlighted Text': highlight.text})
    return send_file(output_file, as_attachment=True)

if __name__ == '__main__':
    os.makedirs('uploads', exist_ok=True)
    os.makedirs('outputs', exist_ok=True)
    with app.app_context():
        db.create_all()  # Create the database tables
    app.run(debug=True)
